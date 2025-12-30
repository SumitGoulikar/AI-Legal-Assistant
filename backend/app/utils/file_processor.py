# backend/app/utils/file_processor.py
"""
File Processing Utilities
=========================
Handles text extraction from various file formats.

Supported formats:
- PDF: Using PyMuPDF (fitz)
- DOCX: Using python-docx
- TXT: Direct reading

Also handles:
- File validation
- Size checking
- Extension verification
"""

import os
import uuid
from typing import Tuple, Optional, List, Dict, Any
from pathlib import Path
import fitz  # PyMuPDF
from docx import Document as DocxDocument
import aiofiles
import aiofiles.os

from app.config import settings
from app.core.exceptions import FileProcessingError, ValidationError


# ============================================
# FILE VALIDATION
# ============================================

def validate_file_extension(filename: str) -> str:
    """
    Validate and return the file extension.
    
    Args:
        filename: Original filename
        
    Returns:
        Lowercase extension without dot (e.g., 'pdf')
        
    Raises:
        ValidationError: If extension is not supported
    """
    ext = Path(filename).suffix.lower()
    
    if ext not in settings.ALLOWED_EXTENSIONS:
        raise ValidationError(
            f"File type '{ext}' is not supported. "
            f"Allowed types: {', '.join(settings.ALLOWED_EXTENSIONS)}"
        )
    
    return ext[1:]  # Remove the dot


def validate_file_size(file_size: int) -> None:
    """
    Validate file size is within limits.
    
    Args:
        file_size: Size in bytes
        
    Raises:
        ValidationError: If file is too large
    """
    if file_size > settings.MAX_FILE_SIZE:
        max_mb = settings.MAX_FILE_SIZE / (1024 * 1024)
        raise ValidationError(
            f"File size exceeds maximum limit of {max_mb:.1f} MB"
        )


def generate_unique_filename(original_name: str) -> str:
    """
    Generate a unique filename for storage.
    
    Args:
        original_name: Original filename
        
    Returns:
        UUID-based unique filename with original extension
    """
    ext = Path(original_name).suffix.lower()
    unique_id = uuid.uuid4().hex[:16]
    return f"{unique_id}{ext}"


def format_file_size(size_bytes: int) -> str:
    """
    Convert bytes to human-readable format.
    
    Args:
        size_bytes: Size in bytes
        
    Returns:
        Formatted string (e.g., "1.5 MB")
    """
    for unit in ['B', 'KB', 'MB', 'GB']:
        if size_bytes < 1024:
            return f"{size_bytes:.1f} {unit}"
        size_bytes /= 1024
    return f"{size_bytes:.1f} TB"


# ============================================
# FILE STORAGE
# ============================================

async def save_upload_file(
    file_content: bytes,
    user_id: str,
    original_filename: str
) -> Tuple[str, str]:
    """
    Save uploaded file to disk.
    
    Args:
        file_content: Raw file bytes
        user_id: User's ID for organizing files
        original_filename: Original filename for extension
        
    Returns:
        Tuple of (stored_filename, full_file_path)
    """
    # Create user directory if it doesn't exist
    user_dir = Path(settings.UPLOAD_DIR) / user_id
    await aiofiles.os.makedirs(user_dir, exist_ok=True)
    
    # Generate unique filename
    stored_filename = generate_unique_filename(original_filename)
    file_path = user_dir / stored_filename
    
    # Write file
    async with aiofiles.open(file_path, 'wb') as f:
        await f.write(file_content)
    
    return stored_filename, str(file_path)


async def delete_file(file_path: str) -> bool:
    """
    Delete a file from disk.
    
    Args:
        file_path: Path to file
        
    Returns:
        True if deleted, False if not found
    """
    try:
        if await aiofiles.os.path.exists(file_path):
            await aiofiles.os.remove(file_path)
            return True
        return False
    except Exception as e:
        print(f"Error deleting file {file_path}: {e}")
        return False


async def get_file_content(file_path: str) -> bytes:
    """
    Read file content from disk.
    
    Args:
        file_path: Path to file
        
    Returns:
        File content as bytes
    """
    async with aiofiles.open(file_path, 'rb') as f:
        return await f.read()


# ============================================
# TEXT EXTRACTION
# ============================================

class ExtractionResult:
    """Result of text extraction from a document."""
    
    def __init__(self):
        self.full_text: str = ""
        self.pages: List[Dict[str, Any]] = []
        self.page_count: int = 0
        self.word_count: int = 0
        self.metadata: Dict[str, Any] = {}
        self.error: Optional[str] = None
    
    def to_dict(self) -> dict:
        return {
            "full_text": self.full_text,
            "pages": self.pages,
            "page_count": self.page_count,
            "word_count": self.word_count,
            "metadata": self.metadata,
            "error": self.error,
        }


def extract_text_from_pdf(file_path: str) -> ExtractionResult:
    """
    Extract text from a PDF file.
    
    Args:
        file_path: Path to PDF file
        
    Returns:
        ExtractionResult with text and metadata
    """
    result = ExtractionResult()
    
    try:
        # Open PDF with PyMuPDF
        doc = fitz.open(file_path)
        
        result.page_count = len(doc)
        result.metadata = {
            "title": doc.metadata.get("title", ""),
            "author": doc.metadata.get("author", ""),
            "subject": doc.metadata.get("subject", ""),
            "creator": doc.metadata.get("creator", ""),
            "creation_date": doc.metadata.get("creationDate", ""),
        }
        
        all_text = []
        
        for page_num, page in enumerate(doc, start=1):
            page_text = page.get_text("text")
            
            result.pages.append({
                "page_num": page_num,
                "text": page_text,
                "char_count": len(page_text),
            })
            
            all_text.append(page_text)
        
        result.full_text = "\n\n".join(all_text)
        result.word_count = len(result.full_text.split())
        
        doc.close()
        
    except Exception as e:
        result.error = f"PDF extraction error: {str(e)}"
    
    return result


def extract_text_from_docx(file_path: str) -> ExtractionResult:
    """
    Extract text from a DOCX file.
    
    Args:
        file_path: Path to DOCX file
        
    Returns:
        ExtractionResult with text and metadata
    """
    result = ExtractionResult()
    
    try:
        doc = DocxDocument(file_path)
        
        # Extract core properties (metadata)
        core_props = doc.core_properties
        result.metadata = {
            "title": core_props.title or "",
            "author": core_props.author or "",
            "subject": core_props.subject or "",
            "created": str(core_props.created) if core_props.created else "",
            "modified": str(core_props.modified) if core_props.modified else "",
        }
        
        # Extract paragraphs
        paragraphs = []
        for para in doc.paragraphs:
            if para.text.strip():
                paragraphs.append(para.text)
        
        # Extract text from tables
        table_texts = []
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text for cell in row.cells)
                if row_text.strip():
                    table_texts.append(row_text)
        
        # Combine all text
        result.full_text = "\n\n".join(paragraphs)
        if table_texts:
            result.full_text += "\n\n[Tables]\n" + "\n".join(table_texts)
        
        result.word_count = len(result.full_text.split())
        result.page_count = 1  # DOCX doesn't have strict pages
        
        # Create a single "page" for consistency
        result.pages = [{
            "page_num": 1,
            "text": result.full_text,
            "char_count": len(result.full_text),
        }]
        
    except Exception as e:
        result.error = f"DOCX extraction error: {str(e)}"
    
    return result


def extract_text_from_txt(file_path: str) -> ExtractionResult:
    """
    Extract text from a TXT file.
    
    Args:
        file_path: Path to TXT file
        
    Returns:
        ExtractionResult with text
    """
    result = ExtractionResult()
    
    try:
        # Try different encodings
        encodings = ['utf-8', 'latin-1', 'cp1252']
        
        for encoding in encodings:
            try:
                with open(file_path, 'r', encoding=encoding) as f:
                    result.full_text = f.read()
                break
            except UnicodeDecodeError:
                continue
        
        if not result.full_text:
            # Fallback: read as bytes and decode with errors='replace'
            with open(file_path, 'rb') as f:
                content = f.read()
                result.full_text = content.decode('utf-8', errors='replace')
        
        result.word_count = len(result.full_text.split())
        result.page_count = 1
        
        result.pages = [{
            "page_num": 1,
            "text": result.full_text,
            "char_count": len(result.full_text),
        }]
        
    except Exception as e:
        result.error = f"TXT extraction error: {str(e)}"
    
    return result


def extract_text(file_path: str, file_type: str) -> ExtractionResult:
    """
    Extract text from a file based on its type.
    
    Args:
        file_path: Path to file
        file_type: File extension (pdf, docx, txt)
        
    Returns:
        ExtractionResult with text and metadata
    """
    file_type = file_type.lower()
    
    if file_type == "pdf":
        return extract_text_from_pdf(file_path)
    elif file_type == "docx":
        return extract_text_from_docx(file_path)
    elif file_type == "txt":
        return extract_text_from_txt(file_path)
    else:
        result = ExtractionResult()
        result.error = f"Unsupported file type: {file_type}"
        return result


# ============================================
# TEXT CHUNKING (for RAG)
# ============================================

def chunk_text(
    text: str,
    chunk_size: int = 500,
    chunk_overlap: int = 50
) -> List[Dict[str, Any]]:
    """
    Split text into overlapping chunks for embedding.
    
    Args:
        text: Full text to chunk
        chunk_size: Target size of each chunk in characters
        chunk_overlap: Number of overlapping characters between chunks
        
    Returns:
        List of chunk dictionaries with text and metadata
    """
    if not text or not text.strip():
        return []
    
    # Clean the text
    text = text.strip()
    
    # Split into sentences (simple approach)
    # In production, use NLTK or spaCy for better sentence splitting
    sentences = []
    for para in text.split('\n'):
        para = para.strip()
        if para:
            # Split by common sentence endings
            import re
            para_sentences = re.split(r'(?<=[.!?])\s+', para)
            sentences.extend([s.strip() for s in para_sentences if s.strip()])
    
    chunks = []
    current_chunk = []
    current_size = 0
    chunk_index = 0
    
    for sentence in sentences:
        sentence_len = len(sentence)
        
        # If adding this sentence exceeds chunk size, save current chunk
        if current_size + sentence_len > chunk_size and current_chunk:
            chunk_text = ' '.join(current_chunk)
            chunks.append({
                "chunk_index": chunk_index,
                "content": chunk_text,
                "char_count": len(chunk_text),
                "start_char": sum(len(c["content"]) for c in chunks),
            })
            chunk_index += 1
            
            # Keep some sentences for overlap
            overlap_size = 0
            overlap_sentences = []
            for s in reversed(current_chunk):
                if overlap_size + len(s) <= chunk_overlap:
                    overlap_sentences.insert(0, s)
                    overlap_size += len(s)
                else:
                    break
            
            current_chunk = overlap_sentences
            current_size = overlap_size
        
        current_chunk.append(sentence)
        current_size += sentence_len
    
    # Don't forget the last chunk
    if current_chunk:
        chunk_text = ' '.join(current_chunk)
        chunks.append({
            "chunk_index": chunk_index,
            "content": chunk_text,
            "char_count": len(chunk_text),
            "start_char": sum(len(c["content"]) for c in chunks[:-1]) if chunks else 0,
        })
    
    return chunks


def assign_pages_to_chunks(
    chunks: List[Dict[str, Any]],
    pages: List[Dict[str, Any]]
) -> List[Dict[str, Any]]:
    """
    Assign page numbers to chunks based on character positions.
    
    Args:
        chunks: List of text chunks
        pages: List of page data with text
        
    Returns:
        Chunks with page information added
    """
    if not pages or len(pages) <= 1:
        # Single page or no page info
        for chunk in chunks:
            chunk["start_page"] = 1
            chunk["end_page"] = 1
        return chunks
    
    # Calculate cumulative character positions for each page
    page_positions = []
    cumulative = 0
    for page in pages:
        page_positions.append({
            "page_num": page["page_num"],
            "start": cumulative,
            "end": cumulative + page["char_count"]
        })
        cumulative += page["char_count"]
    
    # Assign pages to chunks
    for chunk in chunks:
        chunk_start = chunk.get("start_char", 0)
        chunk_end = chunk_start + chunk["char_count"]
        
        start_page = 1
        end_page = 1
        
        for pp in page_positions:
            if pp["start"] <= chunk_start < pp["end"]:
                start_page = pp["page_num"]
            if pp["start"] < chunk_end <= pp["end"]:
                end_page = pp["page_num"]
        
        chunk["start_page"] = start_page
        chunk["end_page"] = end_page
    
    return chunks